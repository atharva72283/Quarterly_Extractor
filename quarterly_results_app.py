import streamlit as st
import pandas as pd
import numpy as np
import io
import re
import os
import json
import base64
import tempfile
import requests
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import datetime

# ─────────────────────────────────────────
# GitHub Config  (private repo – API auth)
# ─────────────────────────────────────────
GITHUB_OWNER    = "atharva72283"
GITHUB_REPO     = "Quarterly_Extractor"
GITHUB_BRANCH   = "main"
GITHUB_API_BASE = "https://api.github.com"

LEFT_LOGO_PATH  = "Q4.png"
RIGHT_LOGO_PATH = "JM_Logo.png"
LOG_FILE_PATH   = "results_log.json"

def get_gh_headers() -> dict:
    """Always build headers at call-time so token is never frozen as None."""
    return {
        "Authorization": f"token {os.getenv('MY_TOKEN', '')}",
        "Accept": "application/vnd.github.v3+json",
    }

# ─────────────────────────────────────────
# Page Config
# ─────────────────────────────────────────
st.set_page_config(
    page_title="JM Financial – Quarterly Results Extractor",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────
# Custom CSS  – JM Financial Gold Theme
# ─────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@300;400;600;700&family=IBM+Plex+Mono:wght@400;600&display=swap');
html, body, [class*="css"] { font-family: 'IBM Plex Sans', sans-serif; }

section[data-testid="stSidebar"] {
    background: #1a1a2e;
    border-right: 2px solid #FCB316;
}
section[data-testid="stSidebar"] * { color: #e8e8e8 !important; }
section[data-testid="stSidebar"] .stTextInput input,
section[data-testid="stSidebar"] .stFileUploader {
    background: #16213e !important;
    border: 1px solid #FCB316 !important;
    color: #fff !important;
    border-radius: 4px;
}

.main .block-container { padding-top: 2rem; max-width: 1400px; }

.jm-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
    border-left: 5px solid #FCB316;
    padding: 1.2rem 1.8rem; border-radius: 6px; margin-bottom: 1.5rem;
    display: flex; align-items: center; gap: 1rem;
}
.jm-header h1 { color: #FCB316; font-size: 1.6rem; font-weight: 700; margin: 0; }
.jm-header p  { color: #aaa; font-size: 0.85rem; margin: 0; }

.step-badge {
    display: inline-block; background: #FCB316; color: #1a1a2e;
    font-weight: 700; font-size: 0.75rem; padding: 2px 10px;
    border-radius: 20px; margin-bottom: 0.4rem;
    font-family: 'IBM Plex Mono', monospace; letter-spacing: 0.05em;
}

div.stButton > button {
    background: #FCB316; color: #1a1a2e; font-weight: 700;
    border: none; border-radius: 4px; padding: 0.5rem 1.5rem;
    font-family: 'IBM Plex Sans', sans-serif; letter-spacing: 0.04em; transition: all 0.2s;
}
div.stButton > button:hover {
    background: #e5a110; transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(252,179,22,0.35);
}

.status-ok   { background:#e8f5e9; border-left:4px solid #4caf50; padding:0.6rem 1rem; border-radius:4px; font-size:0.9rem; margin:0.5rem 0; }
.status-warn { background:#fff8e1; border-left:4px solid #FCB316;  padding:0.6rem 1rem; border-radius:4px; font-size:0.9rem; margin:0.5rem 0; }

.stDataEditor { border: 2px solid #FCB316; border-radius: 6px; }

div.stDownloadButton > button {
    background:#1a1a2e; color:#FCB316; border:2px solid #FCB316; font-weight:700; border-radius:4px;
}
div.stDownloadButton > button:hover { background:#FCB316; color:#1a1a2e; }

.metric-row  { display:flex; gap:0.8rem; flex-wrap:wrap; margin:0.8rem 0; }
.metric-card {
    background:#1a1a2e; border-radius:6px; padding:0.7rem 1.2rem;
    flex:1; min-width:130px; border-bottom:3px solid #FCB316;
}
.metric-card .label { color:#aaa; font-size:0.7rem; text-transform:uppercase; letter-spacing:0.08em; }
.metric-card .value { color:#FCB316; font-size:1.3rem; font-weight:700; font-family:'IBM Plex Mono',monospace; }
.metric-card .sub   { color:#888; font-size:0.72rem; margin-top:1px; }

.ai-summary-box {
    background:#f0f4ff; border:1px solid #c5d3f0; border-left:4px solid #3a5fc8;
    border-radius:6px; padding:1rem 1.4rem; margin:0.8rem 0;
    font-size:0.9rem; line-height:1.6;
}
.ai-summary-box h4 { color:#1a1a2e; margin-bottom:0.5rem; font-size:0.95rem; }

/* Results Log table */
.log-table { width:100%; border-collapse:collapse; font-size:0.85rem; margin-top:0.5rem; }
.log-table th {
    background:#FCB316; color:#1a1a2e; font-weight:700;
    padding:6px 10px; text-align:left; border:1px solid #e0c060;
}
.log-table td { padding:5px 10px; border:1px solid #ddd; background:#fff; }
.log-table tr:nth-child(even) td { background:#fafafa; }
.tag-positive {
    background:#d4edda; color:#155724; font-weight:700;
    padding:2px 8px; border-radius:10px; font-size:0.78rem;
}
.tag-negative {
    background:#f8d7da; color:#721c24; font-weight:700;
    padding:2px 8px; border-radius:10px; font-size:0.78rem;
}
.company-positive { color:#155724; font-weight:700; }
.company-negative { color:#721c24; font-weight:700; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# Header
# ─────────────────────────────────────────
st.markdown("""
<div class="jm-header">
    <div>
        <h1>📊 Quarterly Results Extractor</h1>
        <p>JM Financial · Quarterly Result Extractor · AI Integrated</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# GitHub helpers
# ─────────────────────────────────────────
@st.cache_data(show_spinner=False)
def fetch_github_image(file_path: str) -> bytes | None:
    """
    Fetch an image from a PRIVATE GitHub repo using the Contents API.
    The API returns the file as base64 in JSON — works for private repos
    unlike raw.githubusercontent.com which 403s without a token in the URL.
    Cache is keyed on file_path; cleared after extraction to force refresh.
    """
    url = f"{GITHUB_API_BASE}/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{file_path}"
    try:
        resp = requests.get(url, headers=get_gh_headers(), timeout=15)
        if resp.status_code == 200:
            data = resp.json()
            # GitHub returns content with newlines – strip them before decoding
            b64_content = data.get("content", "").replace("\n", "")
            if b64_content:
                return base64.b64decode(b64_content)
        else:
            st.warning(f"GitHub image fetch failed for '{file_path}' "
                       f"(HTTP {resp.status_code}). Check filename & MY_TOKEN.")
    except Exception as e:
        st.warning(f"Could not fetch '{file_path}' from GitHub: {e}")
    return None


def fetch_results_log() -> list:
    """
    Read results_log.json from the repo via Contents API.
    NOT cached — must be live every render so sidebar shows fresh data
    immediately after a push without needing a manual page refresh.
    """
    url = f"{GITHUB_API_BASE}/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{LOG_FILE_PATH}"
    try:
        resp = requests.get(url, headers=get_gh_headers(), timeout=15)
        if resp.status_code == 200:
            b64_content = resp.json().get("content", "").replace("\n", "")
            raw = base64.b64decode(b64_content).decode("utf-8")
            return json.loads(raw)
    except Exception:
        pass
    return []


def push_results_log(log_entries: list) -> bool:
    """
    Commit the updated results_log.json back to the repo.
    Creates the file if it doesn't exist, updates (with SHA) if it does.
    Returns True on success.
    """
    url     = f"{GITHUB_API_BASE}/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{LOG_FILE_PATH}"
    content = base64.b64encode(
        json.dumps(log_entries, indent=2, ensure_ascii=False).encode("utf-8")
    ).decode("utf-8")

    # Must supply SHA when updating an existing file
    sha = None
    try:
        resp = requests.get(url, headers=get_gh_headers(), timeout=15)
        if resp.status_code == 200:
            sha = resp.json().get("sha")
    except Exception:
        pass

    payload = {
        "message": "results_log: add/update entry",
        "content": content,
        "branch":  GITHUB_BRANCH,
    }
    if sha:
        payload["sha"] = sha

    try:
        resp = requests.put(url, headers=get_gh_headers(), json=payload, timeout=20)
        return resp.status_code in (200, 201)
    except Exception as e:
        st.warning(f"Could not push log to GitHub: {e}")
        return False


# ─────────────────────────────────────────
# Pre-load images from GitHub (once per session via cache)
# Uses Contents API so works on private repos
# ─────────────────────────────────────────
left_logo_bytes  = fetch_github_image(LEFT_LOGO_PATH)
right_logo_bytes = fetch_github_image(RIGHT_LOGO_PATH)

# ─────────────────────────────────────────
# Sidebar – Credentials & Settings
# ─────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🔑 API Configuration")
    api_key = st.text_input(
        "Mistral API Key",
        type="password",
        placeholder="Enter your Mistral API key...",
        help="Your Mistral API key from console.mistral.ai"
    )

    st.markdown("---")
    st.markdown("### 📋 Report Settings")

    company_name = st.text_input(
        "Company Name",
        placeholder="e.g. Reliance Industries Ltd.",
        value=""
    )

    page_number = st.number_input(
        "PDF Page to Analyze",
        min_value=1, max_value=500, value=1, step=1,
        help="The page number (1-indexed) containing the financial table"
    )

    # ── Permanent Results Log ───────────────
    st.markdown("---")
    st.markdown("### 📁 Results Log")
    st.caption("Stored permanently in GitHub repo")

    log_entries = fetch_results_log()

    if log_entries:
        st.markdown(
            f"<small style='color:#aaa'>{len(log_entries)} entr{'y' if len(log_entries)==1 else 'ies'} recorded</small>",
            unsafe_allow_html=True
        )
        # Build HTML table
        rows_html = ""
        for entry in reversed(log_entries[-20:]):   # show latest 20
            sentiment = entry.get("sentiment", "").lower()
            company_class = "company-positive" if sentiment == "positive" else "company-negative"
            tag_class     = "tag-positive"      if sentiment == "positive" else "tag-negative"
            tag_label     = "▲ Positive"         if sentiment == "positive" else "▼ Negative"
            rows_html += f"""
            <tr>
              <td class="{company_class}">{entry.get('company','')}</td>
              <td>{entry.get('revenue','–')}</td>
              <td>{entry.get('pat','–')}</td>
              <td><span class="{tag_class}">{tag_label}</span></td>
              <td style="color:#888;font-size:0.75rem">{entry.get('date','')}</td>
            </tr>"""

        st.markdown(f"""
        <table class="log-table">
          <thead>
            <tr>
              <th>Company</th><th>Revenue</th><th>PAT</th><th>Result</th><th>Date</th>
            </tr>
          </thead>
          <tbody>{rows_html}</tbody>
        </table>
        """, unsafe_allow_html=True)
    else:
        st.markdown("<small style='color:#888'>No entries yet. Generate a report to start logging.</small>", unsafe_allow_html=True)

    st.markdown("---")
    st.caption("v3.0 · JM Financial Internal Tool · Mistral AI")

# ─────────────────────────────────────────
# Session-state init
# ─────────────────────────────────────────
for key in ["df_extracted", "mistral_raw", "step", "ai_summary", "pdf_text_full"]:
    if key not in st.session_state:
        st.session_state[key] = None
if "step" not in st.session_state or st.session_state.step is None:
    st.session_state.step = 1

# ─────────────────────────────────────────
# Constants & API config
# ─────────────────────────────────────────
COLS_NUMERIC         = ['Q4FY2026', 'Q3FY2026', 'Q4FY2025', 'FY2026', 'FY2025']
MISTRAL_VISION_MODEL = "pixtral-12b-2409"
MISTRAL_TEXT_MODEL   = "mistral-large-latest"
MISTRAL_API_BASE     = "https://api.mistral.ai/v1"


# ─────────────────────────────────────────
# Core helpers
# ─────────────────────────────────────────
def extract_pdf_page_as_image(pdf_bytes: bytes, page_idx: int) -> Image.Image:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if page_idx >= len(doc):
        st.error(f"PDF only has {len(doc)} pages.")
        st.stop()
    page = doc[page_idx]
    pix  = page.get_pixmap(dpi=200)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def extract_pdf_full_text(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def pil_to_b64(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def mistral_headers(key: str) -> dict:
    return {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}


def call_mistral_vision(api_key: str, img: Image.Image) -> str:
    prompt = """Extract the financial data from the provided image into a clean CSV format.
Columns MUST exactly match: Particulars, Q4FY2026, Q3FY2026, Q4FY2025, FY2026, FY2025.

Extract or calculate the following rows exactly in this order:
1.  Revenue from Operations
2.  Other Incomes
3.  Total Income         -> (Calculate: Revenue from Operations + Other Incomes)
4.  Total Expenditure    -> (Calculate: Total Expenses - Interest - Depreciation)
5.  EBITDA               -> (Calculate: Total Income - Total Expenditure)
6.  Interest             -> (Finance cost)
7.  Depreciation         -> (Depreciation and amortisation)
8.  PBT                  -> (Profit Before Tax)
9.  Tax                  -> (Total tax expense)
10. Provisions & Contingencies
11. Share of Profit of Associate/JV
12. Exceptional Items
13. PAT                  -> (Profit After Tax / Net Profit for the period)
14. EPS                  -> (Earnings Per Share - Basic/Diluted)

IMPORTANT: Return ONLY the CSV block. No intro, no outro, no explanations.
Use 0 for missing values. Clean numbers of all currency symbols and commas.
Indicate negative numbers with a minus sign (e.g., -100)."""

    payload = {
        "model": MISTRAL_VISION_MODEL,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": f"data:image/png;base64,{pil_to_b64(img)}"},
                {"type": "text",      "text": prompt}
            ]
        }],
        "max_tokens": 2000
    }
    resp = requests.post(f"{MISTRAL_API_BASE}/chat/completions",
                         json=payload, headers=mistral_headers(api_key), timeout=90)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def call_mistral_ai_summary(api_key: str, pdf_text: str, company: str) -> str:
    system_prompt = (
        "You are a risk research analyst. This is the quarterly result of a company. "
        "Search for any key matters highlighted by company in this quarter such as acquisitions, "
        "penalty, litigations etc and give outcome in bullet form. "
        "Further give your 2 liner view on the financial results of the company."
    )
    user_prompt = (
        f"Company: {company}\n\n"
        f"Quarterly Results Document Text:\n{pdf_text[:15000]}\n\n"
        "Please provide:\n"
        "1. Key matters (acquisitions, penalties, litigations, one-time items, management commentary highlights) in bullet points\n"
        "2. Your 2-line analyst view on the financial results"
    )
    payload = {
        "model": MISTRAL_TEXT_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt}
        ],
        "max_tokens": 1000
    }
    resp = requests.post(f"{MISTRAL_API_BASE}/chat/completions",
                         json=payload, headers=mistral_headers(api_key), timeout=90)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def call_mistral_sentiment(api_key: str, pat_yoy: float, pat_qoq: float,
                            revenue_yoy: float, summary: str) -> str:
    """
    Ask Mistral to classify the result as 'positive' or 'negative'
    based on financials + AI summary context.
    Returns exactly 'positive' or 'negative'.
    """
    prompt = (
        f"You are a financial analyst. Based on the following quarterly result data, "
        f"classify the overall result as either 'positive' or 'negative'. "
        f"Reply with ONLY one word: positive or negative.\n\n"
        f"PAT YoY change: {pat_yoy:+.1f}%\n"
        f"PAT QoQ change: {pat_qoq:+.1f}%\n"
        f"Revenue YoY change: {revenue_yoy:+.1f}%\n"
        f"Analyst commentary excerpt: {summary[:500]}"
    )
    payload = {
        "model": MISTRAL_TEXT_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 5
    }
    resp = requests.post(f"{MISTRAL_API_BASE}/chat/completions",
                         json=payload, headers=mistral_headers(api_key), timeout=30)
    resp.raise_for_status()
    raw = resp.json()["choices"][0]["message"]["content"].strip().lower()
    return "positive" if "positive" in raw else "negative"


def parse_mistral_csv(raw: str) -> pd.DataFrame:
    csv_match = re.search(r"(Particulars,.*)", raw, re.DOTALL)
    clean_csv = (
        csv_match.group(1).replace("```", "").strip()
        if csv_match
        else raw.replace("```csv", "").replace("```", "").strip()
    )
    df = pd.read_csv(io.StringIO(clean_csv), skipinitialspace=True)
    for col in COLS_NUMERIC:
        if col in df.columns:
            df[col] = pd.to_numeric(
                df[col].astype(str)
                    .str.replace(",", "", regex=False)
                    .str.replace(r"\(", "-", regex=True)
                    .str.replace(r"\)", "",  regex=True),
                errors="coerce"
            ).fillna(0)
    df["QoQ%"]    = ((df["Q4FY2026"] - df["Q3FY2026"]) / df["Q3FY2026"].replace(0, np.nan) * 100).fillna(0).round(0).astype(int)
    df["YoY%"]    = ((df["Q4FY2026"] - df["Q4FY2025"]) / df["Q4FY2025"].replace(0, np.nan) * 100).fillna(0).round(0).astype(int)
    df["% Change"]= ((df["FY2026"]   - df["FY2025"])   / df["FY2025"].replace(0, np.nan)   * 100).fillna(0).round(0).astype(int)
    return df[["Particulars","Q4FY2026","Q3FY2026","Q4FY2025","QoQ%","YoY%","FY2026","FY2025","% Change"]]


def get_row(df: pd.DataFrame, label: str) -> pd.Series:
    row = df[df["Particulars"].str.strip() == label]
    return row.iloc[0] if not row.empty else pd.Series({c: 0 for c in df.columns})


def build_template_df(df: pd.DataFrame) -> pd.DataFrame:
    net_sales = get_row(df, "Revenue from Operations")
    other_inc = get_row(df, "Other Incomes")
    tot_inc   = get_row(df, "Total Income")
    tot_exp   = get_row(df, "Total Expenditure")
    ebitda    = get_row(df, "EBITDA")
    interest  = get_row(df, "Interest")
    dep       = get_row(df, "Depreciation")
    pbt       = get_row(df, "PBT")
    tax       = get_row(df, "Tax")
    prov      = get_row(df, "Provisions & Contingencies")
    exc       = get_row(df, "Exceptional Items")
    pat       = get_row(df, "PAT")
    eps       = get_row(df, "EPS")

    template_rows = [
        "Net Sales","Other Income","Total Income","",
        "Total Expenditure (Ex Int & Dep)","EBIDTA","Interest",
        "Dep","Profit before tax","Tax","Provisions & Contingencies",
        "Exceptional Item","Reported Profit","EBIDTA Margin %","PAT %","EPS (Diluted)"
    ]
    all_cols = ["Q4FY2026","Q3FY2026","Q4FY2025","QoQ%","YoY%","FY2026","FY2025","% Change"]
    mapped = {c: [] for c in all_cols}

    for col in all_cols:
        ns = net_sales.get(col, 0); oi = other_inc.get(col, 0)
        ti = tot_inc.get(col, 0);   te = tot_exp.get(col, 0)
        eb = ebitda.get(col, 0);    ir = interest.get(col, 0)
        dp = dep.get(col, 0);       pb = pbt.get(col, 0)
        tx = tax.get(col, 0);       pv = prov.get(col, 0)
        ex = exc.get(col, 0);       pt = pat.get(col, 0)
        ep = eps.get(col, 0)
        eb_m = (eb / ns * 100) if ns != 0 else 0
        pt_m = (pt / ns * 100) if ns != 0 else 0
        mapped[col] = [ns,oi,ti,"",te,eb,ir,dp,pb,tx,pv,ex,pt,eb_m,pt_m,ep]

    tdf = pd.DataFrame({"Particulars": template_rows})
    for col in all_cols:
        tdf[col] = mapped[col]
    return tdf


def generate_paragraphs(df: pd.DataFrame, company: str) -> list:
    ti  = get_row(df, "Total Income")
    eb  = get_row(df, "EBITDA")
    pat = get_row(df, "PAT")
    eps = get_row(df, "EPS")
    tdf = build_template_df(df)

    def gv(row_label, col):
        r = tdf[tdf["Particulars"] == row_label]
        return float(r.iloc[0][col]) if not r.empty else 0

    ebitda_margin_q4 = gv("EBIDTA Margin %", "Q4FY2026")
    ebitda_margin_q3 = gv("EBIDTA Margin %", "Q3FY2026")
    pat_margin_q4    = gv("PAT %",           "Q4FY2026")
    pat_margin_q3    = gv("PAT %",           "Q3FY2026")
    ebitda_bps = (ebitda_margin_q4 - ebitda_margin_q3) * 100
    pat_bps    = (pat_margin_q4    - pat_margin_q3)    * 100

    fn  = lambda v: f"{float(v):,.2f}"
    fp  = lambda v: f"{abs(float(v)):.0f}"
    ud  = lambda v: "Increased" if float(v) >= 0 else "Decreased"
    udl = lambda v: "increased" if float(v) >= 0 else "decreased"

    return [
        (f"{company} reported net profit of INR {fn(pat['Q4FY2026'])} Crs. in Q4FY2026, "
         f"{ud(pat['QoQ%'])} by ~{fp(pat['QoQ%'])}% QoQ & {ud(pat['YoY%'])} by ~{fp(pat['YoY%'])}% YoY "
         f"(Margin percentage {fn(pat_margin_q4)}%)."),
        (f"Company Reported Total Income of INR {fn(ti['Q4FY2026'])} Crs. in Q4FY2026 "
         f"as compared to INR {fn(ti['Q3FY2026'])} Crs. in previous quarter Q3FY2026. "
         f"Total Income has {udl(ti['QoQ%'])} by {fp(ti['QoQ%'])}% QoQ. "
         f"(Total Income of INR {fn(ti['Q4FY2025'])} Crs. in Q4FY2025 {udl(ti['YoY%'])} by {fp(ti['YoY%'])}% YoY) "
         f"For FY2026, Company has Reported Total Income of INR {fn(ti['FY2026'])} Crs. "
         f"as against INR {fn(ti['FY2025'])} Crs. in FY2025 {udl(ti['% Change'])} by ~{fp(ti['% Change'])}%."),
        (f"Company Reported EBITDA of INR {fn(eb['Q4FY2026'])} Crs. for Q4FY2026 "
         f"as compared to INR {fn(eb['Q3FY2026'])} in previous quarter Q3FY2026. "
         f"EBITDA has {udl(eb['QoQ%'])} by {fp(eb['QoQ%'])}% QoQ. "
         f"(EBIDTA of INR {fn(eb['Q4FY2025'])} Crs. in Q4FY2025 {udl(eb['YoY%'])} by {fp(eb['YoY%'])}% YoY) "
         f"EBIDTA margin stood at {fn(ebitda_margin_q4)}% for the quarter "
         f"{udl(ebitda_bps)} by {int(abs(ebitda_bps))} bps. "
         f"For FY2026, Company has Reported EBIDTA of INR {fn(eb['FY2026'])} Crs. "
         f"as against INR {fn(eb['FY2025'])} Crs. in FY2025 {ud(eb['% Change'])} by ~{fp(eb['% Change'])}%."),
        (f"Company Reported Profit of INR {fn(pat['Q4FY2026'])} Crs. for Q4FY2026 "
         f"as compared to INR {fn(pat['Q3FY2026'])} Crs. in previous quarter Q3FY2026. "
         f"Profit has {udl(pat['QoQ%'])} by {fp(pat['QoQ%'])}% QoQ. "
         f"(Profit of INR {fn(pat['Q4FY2025'])} Crs. in Q4FY2025 {udl(pat['YoY%'])} by {fp(pat['YoY%'])}% YoY.) "
         f"Profit margin stood at {fn(pat_margin_q4)}% for the quarter "
         f"{udl(pat_bps)} by {int(abs(pat_bps))} bps. "
         f"For FY2026, Company has Reported Profit of INR {fn(pat['FY2026'])} Crs. "
         f"as against INR {fn(pat['FY2025'])} Crs. in FY2025 {ud(pat['% Change'])} by ~{fp(pat['% Change'])}%."),
        (f"Company reported EPS of INR {fn(eps['Q4FY2026'])} "
         f"{udl(eps['QoQ%'])} by {fp(eps['QoQ%'])}% QoQ & "
         f"for FY2026 EPS stood at INR {fn(eps['FY2026'])}."),
    ]


def build_table_image(tdf: pd.DataFrame) -> bytes:
    df_img   = tdf.copy()
    HIGHLIGHT = ["Total Income","EBIDTA","Profit before tax","Reported Profit"]

    for i in range(len(df_img)):
        row_label = str(df_img.iloc[i, 0]).strip()
        is_eps    = row_label == "EPS (Diluted)"
        is_blank  = row_label == ""
        for col in df_img.columns:
            if col == "Particulars":
                continue
            val        = tdf.loc[i, col]
            is_pct_col = col in ["QoQ%", "YoY%", "% Change"]
            is_pct_row = row_label in ["EBIDTA Margin %", "PAT %"]
            if is_blank:
                df_img.loc[i, col] = ""
            elif pd.isna(val) or val == 0 or val == "":
                df_img.loc[i, col] = "-"
            else:
                try:
                    vf = float(val)
                    if is_pct_col or is_pct_row:
                        vi = int(round(vf))
                        df_img.loc[i, col] = f"({abs(vi)}%)" if vi < 0 else f"{vi}%"
                    elif is_eps:
                        df_img.loc[i, col] = f"({abs(vf):,.2f})" if vf < 0 else f"{vf:,.2f}"
                    else:
                        vi = int(round(vf))
                        df_img.loc[i, col] = f"({abs(vi):,})" if vi < 0 else f"{vi:,}"
                except Exception:
                    df_img.loc[i, col] = str(val)

    fig, ax = plt.subplots(figsize=(15, df_img.shape[0] * 0.34 + 0.5))
    ax.axis("off")
    custom_widths = [0.27] + [0.09] * (len(df_img.columns) - 1)
    table = ax.table(
        cellText=df_img.values, colLabels=df_img.columns,
        loc="center", cellLoc="right", colWidths=custom_widths,
    )
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1, 1.4)

    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor("#FCB316")
            cell.set_text_props(weight="bold", color="black", ha="center")
        else:
            row_label = str(df_img.iloc[row - 1, 0]).strip()
            if col == 0:
                cell.set_text_props(ha="left")
            if row_label in HIGHLIGHT:
                cell.set_facecolor("#EBEBEB")
                if col == 0:
                    cell.set_text_props(weight="bold")
        cell.set_edgecolor("black")

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight", dpi=220)
    plt.close()
    buf.seek(0)
    return buf.read()


def add_para(doc, text, bold=False, font_name="Segoe UI", size=12,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p   = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold      = bold
    return p


def build_word_doc(company, paragraphs, table_png, ai_summary="",
                   left_img=None, right_img=None) -> bytes:
    doc     = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)

    # ── Header: left = Q4.png, right = JM Logo.png ──────────────────
    header = section.header
    for para in list(header.paragraphs):
        try: para._element.getparent().remove(para._element)
        except Exception: pass

    htbl = header.add_table(rows=1, cols=2, width=Inches(7.5))
    htbl.autofit = False
    htbl.columns[0].width = Inches(3.75)
    htbl.columns[1].width = Inches(3.75)

    lp = htbl.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if left_img:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            tf.write(left_img); tf.flush()
            lp.add_run().add_picture(tf.name, width=Inches(2.5))

    rp = htbl.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if right_img:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            tf.write(right_img); tf.flush()
            rp.add_run().add_picture(tf.name, width=Inches(1.5))

    # ── Page border ──────────────────────────────────────────────────
    sectPr    = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for bn in ["top","left","bottom","right"]:
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "24");   b.set(qn("w:color"), "auto")
        pgBorders.append(b)
    sectPr.append(pgBorders)

    # ── Earnings paragraphs (Segoe UI 12pt) ─────────────────────────
    for i, text in enumerate(paragraphs):
        add_para(doc, text, bold=(i == 0), font_name="Segoe UI", size=12)

    doc.add_paragraph()

    # ── Table image ──────────────────────────────────────────────────
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
        tf.write(table_png); tf.flush()
        doc.add_picture(tf.name, width=Inches(7.5))

    # ── AI Risk Analyst Summary (Segoe UI 12pt) ──────────────────────
    if ai_summary and ai_summary.strip():
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = h.add_run("Risk Analyst Commentary (AI-Generated)")
        hr.font.name = "Segoe UI"; hr.font.size = Pt(12); hr.bold = True
        hr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        for line in ai_summary.split("\n"):
            stripped = line.strip()
            if stripped:
                add_para(doc, stripped, bold=False, font_name="Segoe UI",
                         size=12, align=WD_ALIGN_PARAGRAPH.LEFT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_excel(df: pd.DataFrame, company: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    tdf = build_template_df(df)
    wb  = Workbook(); ws = wb.active; ws.title = "Financials"
    GOLD = "FCB316"; LGREY = "EBEBEB"
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    HIGHLIGHT = {"Total Income","EBIDTA","Profit before tax","Reported Profit"}

    last_col = get_column_letter(len(tdf.columns))
    ws.merge_cells(f"A1:{last_col}1")
    tc = ws["A1"]
    tc.value = f"{company} – Quarterly Result Update  (Amount in Crs)"
    tc.font = Font(name="Calibri", size=11, bold=True)
    tc.alignment = Alignment(horizontal="center", vertical="center")
    tc.border = border

    for ci, col_name in enumerate(tdf.columns, start=1):
        cell = ws.cell(row=2, column=ci, value=col_name)
        cell.font = Font(bold=True, color="000000")
        cell.fill = PatternFill(start_color=GOLD, end_color=GOLD, fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    NUM_FMT = '#,##0_);[Black](#,##0);"-"'
    PCT_FMT = '0"%";[Black]\\(0"%"\\);"-"'

    for ri, row in tdf.iterrows():
        for ci, col_name in enumerate(tdf.columns, start=1):
            val  = row[col_name]
            cell = ws.cell(row=ri+3, column=ci)
            try:   cell.value = float(val) if val != "" else None
            except: cell.value = val
            if col_name != "Particulars":
                is_pct = col_name in ["QoQ%","YoY%","% Change"] or \
                         str(row["Particulars"]).strip() in ["EBIDTA Margin %","PAT %"]
                cell.number_format = PCT_FMT if is_pct else NUM_FMT
                cell.alignment = Alignment(horizontal="right")
            else:
                cell.alignment = Alignment(horizontal="left")
            cell.border = border
            if str(row["Particulars"]).strip() in HIGHLIGHT:
                cell.fill = PatternFill(start_color=LGREY, end_color=LGREY, fill_type="solid")
                if col_name == "Particulars":
                    cell.font = Font(bold=True)

    for ci, col_name in enumerate(tdf.columns, start=1):
        cl = get_column_letter(ci)
        ws.column_dimensions[cl].width = 34 if col_name == "Particulars" else 14

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────
# STEP 1 – PDF Upload & Extraction
# ─────────────────────────────────────────
st.markdown('<div class="step-badge">STEP 1</div>', unsafe_allow_html=True)
st.subheader("Upload PDF & Extract Data")

pdf_file = st.file_uploader(
    "Upload the quarterly results PDF", type=["pdf"],
    help="The PDF containing the P&L / financial results page"
)

col_btn, col_status = st.columns([1, 3])
with col_btn:
    extract_btn = st.button("⚡ Extract with Mistral", use_container_width=True)

if extract_btn:
    if not api_key:
        st.error("Please enter your Mistral API key in the sidebar.")
    elif not pdf_file:
        st.error("Please upload a PDF file.")
    elif not company_name.strip():
        st.error("Please enter the company name in the sidebar.")
    else:
        pdf_bytes = pdf_file.read()

        with st.spinner(f"Extracting page {page_number} from PDF…"):
            img = extract_pdf_page_as_image(pdf_bytes, page_number - 1)

        with st.spinner("Extracting full PDF text for AI summary…"):
            pdf_text = extract_pdf_full_text(pdf_bytes)
            st.session_state.pdf_text_full = pdf_text

        with st.spinner("Sending table image to Mistral Pixtral (vision)…"):
            try:
                raw = call_mistral_vision(api_key, img)
                st.session_state.mistral_raw = raw
            except Exception as e:
                st.error(f"Mistral Vision API error: {e}")
                st.stop()

        with st.spinner("Generating AI Risk Analyst Summary…"):
            try:
                ai_summary = call_mistral_ai_summary(api_key, pdf_text, company_name)
                st.session_state.ai_summary = ai_summary
            except Exception as e:
                st.warning(f"AI summary failed (continuing without it): {e}")
                st.session_state.ai_summary = ""

        with st.spinner("Parsing extracted data…"):
            df = parse_mistral_csv(raw)
            st.session_state.df_extracted = df
            st.session_state.step = 2

        # ── Compute sentiment & update permanent log ─────────────────
        with st.spinner("Classifying result sentiment & updating log…"):
            try:
                pat_row = get_row(df, "PAT")
                rev_row = get_row(df, "Revenue from Operations")
                sentiment = call_mistral_sentiment(
                    api_key,
                    pat_yoy=float(pat_row.get("YoY%", 0)),
                    pat_qoq=float(pat_row.get("QoQ%", 0)),
                    revenue_yoy=float(rev_row.get("YoY%", 0)),
                    summary=st.session_state.ai_summary or ""
                )
            except Exception:
                sentiment = "positive" if float(pat_row.get("YoY%", 0)) >= 0 else "negative"

            new_entry = {
                "company":   company_name.strip(),
                "revenue":   f"₹{float(rev_row.get('Q4FY2026', 0)):,.0f} Cr",
                "pat":       f"₹{float(pat_row.get('Q4FY2026', 0)):,.0f} Cr",
                "sentiment": sentiment,
                "date":      datetime.date.today().strftime("%d %b %Y"),
                "pat_yoy":   f"{float(pat_row.get('YoY%', 0)):+.0f}%",
                "rev_yoy":   f"{float(rev_row.get('YoY%', 0)):+.0f}%",
            }

            # Fetch current log, append, push back
            current_log = fetch_results_log()
            # Upsert: if same company + date exists, replace it
            current_log = [e for e in current_log
                           if not (e.get("company") == new_entry["company"]
                                   and e.get("date") == new_entry["date"])]
            current_log.append(new_entry)
            push_ok = push_results_log(current_log)
            if push_ok:
                st.markdown('<div class="status-ok">✅ Data extracted & results log updated on GitHub.</div>',
                            unsafe_allow_html=True)
            else:
                st.markdown('<div class="status-ok">✅ Data extracted successfully! '
                            '(Log push to GitHub failed – check token permissions.)</div>',
                            unsafe_allow_html=True)
            # Clear image cache so it re-fetches on next session start if needed
            fetch_github_image.clear()

# ─────────────────────────────────────────
# STEP 2 – Review & Edit
# ─────────────────────────────────────────
if st.session_state.df_extracted is not None:
    st.markdown("---")
    st.markdown('<div class="step-badge">STEP 2</div>', unsafe_allow_html=True)
    st.subheader("Review & Edit Extracted Data")
    st.caption("Double-click any numeric cell to correct values.")

    edited_df = st.data_editor(
        st.session_state.df_extracted,
        use_container_width=True,
        num_rows="fixed",
        key="main_editor",
        column_config={
            "Particulars": st.column_config.TextColumn("Particulars", width="medium"),
            "Q4FY2026": st.column_config.NumberColumn("Q4FY2026", format="%.2f"),
            "Q3FY2026": st.column_config.NumberColumn("Q3FY2026", format="%.2f"),
            "Q4FY2025": st.column_config.NumberColumn("Q4FY2025", format="%.2f"),
            "QoQ%":     st.column_config.NumberColumn("QoQ%",     format="%d"),
            "YoY%":     st.column_config.NumberColumn("YoY%",     format="%d"),
            "FY2026":   st.column_config.NumberColumn("FY2026",   format="%.2f"),
            "FY2025":   st.column_config.NumberColumn("FY2025",   format="%.2f"),
            "% Change": st.column_config.NumberColumn("% Change", format="%d"),
        }
    )

    # KPI strip
    try:
        pat_row = get_row(edited_df, "PAT")
        eb_row  = get_row(edited_df, "EBITDA")
        rev_row = get_row(edited_df, "Revenue from Operations")
        eb_mg   = (eb_row["Q4FY2026"] / rev_row["Q4FY2026"] * 100) if rev_row["Q4FY2026"] else 0
        pt_mg   = (pat_row["Q4FY2026"] / rev_row["Q4FY2026"] * 100) if rev_row["Q4FY2026"] else 0
        st.markdown(f"""
        <div class="metric-row">
          <div class="metric-card"><div class="label">Revenue (Q4FY26)</div>
            <div class="value">₹{rev_row['Q4FY2026']:,.0f}</div><div class="sub">Crs.</div></div>
          <div class="metric-card"><div class="label">EBITDA (Q4FY26)</div>
            <div class="value">₹{eb_row['Q4FY2026']:,.0f}</div><div class="sub">{eb_mg:.1f}% margin</div></div>
          <div class="metric-card"><div class="label">PAT (Q4FY26)</div>
            <div class="value">₹{pat_row['Q4FY2026']:,.0f}</div><div class="sub">{pt_mg:.1f}% margin</div></div>
          <div class="metric-card"><div class="label">PAT QoQ</div>
            <div class="value">{pat_row['QoQ%']:+.0f}%</div><div class="sub">vs Q3FY26</div></div>
          <div class="metric-card"><div class="label">PAT YoY</div>
            <div class="value">{pat_row['YoY%']:+.0f}%</div><div class="sub">vs Q4FY25</div></div>
        </div>""", unsafe_allow_html=True)
    except Exception:
        pass

    # AI Summary display
    if st.session_state.ai_summary:
        st.markdown("---")
        st.markdown("#### 🤖 AI Risk Analyst Summary")
        st.markdown(
            f'<div class="ai-summary-box"><h4>Key Matters &amp; Analyst View</h4>'
            f'{st.session_state.ai_summary.replace(chr(10), "<br>")}</div>',
            unsafe_allow_html=True
        )

    # Recalculate growth cols
    for col in COLS_NUMERIC:
        if col not in edited_df.columns:
            edited_df[col] = 0
    edited_df["QoQ%"]    = ((edited_df["Q4FY2026"]-edited_df["Q3FY2026"])/edited_df["Q3FY2026"].replace(0,np.nan)*100).fillna(0).round(0).astype(int)
    edited_df["YoY%"]    = ((edited_df["Q4FY2026"]-edited_df["Q4FY2025"])/edited_df["Q4FY2025"].replace(0,np.nan)*100).fillna(0).round(0).astype(int)
    edited_df["% Change"]= ((edited_df["FY2026"]  -edited_df["FY2025"])  /edited_df["FY2025"].replace(0,np.nan)*100).fillna(0).round(0).astype(int)

    # ─────────────────────────────────────────
    # STEP 3 – Generate Outputs
    # ─────────────────────────────────────────
    st.markdown("---")
    st.markdown('<div class="step-badge">STEP 3</div>', unsafe_allow_html=True)
    st.subheader("Generate & Download Outputs")

    tab_excel, tab_word = st.tabs(["📗  Excel Output", "📝  Word Report"])

    with tab_excel:
        st.markdown("Generates a styled Excel file matching the JM Financial quarterly template.")
        if st.button("Build Excel", key="build_excel"):
            with st.spinner("Building Excel…"):
                xl_bytes = build_excel(edited_df, company_name)
            st.download_button(
                label="⬇️  Download Excel", data=xl_bytes,
                file_name=f"{company_name} Q4_Result.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    with tab_word:
        st.markdown(
            "Generates the styled Word document with earnings commentary, "
            "coloured table, and AI Risk Analyst Summary. "
            "Header images are pulled automatically from GitHub."
        )

        paras = generate_paragraphs(edited_df, company_name)
        st.markdown("**Generated Earnings Summary**")
        for i, p in enumerate(paras):
            st.text_area(f"Paragraph {i+1}", value=p, height=90, key=f"para_{i}")

        if st.session_state.ai_summary:
            st.markdown("**AI Risk Analyst Summary** *(editable)*")
            ai_summary_edited = st.text_area(
                "AI Summary (appears below table in Word)",
                value=st.session_state.ai_summary,
                height=200, key="ai_summary_edit"
            )
        else:
            ai_summary_edited = ""

        st.caption("Font: Segoe UI 12pt throughout. Header images loaded from GitHub automatically.")

        if st.button("Build Word Document", key="build_word"):
            final_paras      = [st.session_state.get(f"para_{i}", p) for i, p in enumerate(paras)]
            final_ai_summary = st.session_state.get("ai_summary_edit", ai_summary_edited)

            with st.spinner("Rendering table image…"):
                tdf       = build_template_df(edited_df)
                table_png = build_table_image(tdf)

            with st.spinner("Building Word document…"):
                docx_bytes = build_word_doc(
                    company_name, final_paras, table_png,
                    ai_summary=final_ai_summary,
                    left_img=left_logo_bytes,
                    right_img=right_logo_bytes,
                )

            st.download_button(
                label="⬇️  Download Word Document", data=docx_bytes,
                file_name=f"{company_name} Q4_Result.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# ─────────────────────────────────────────
# Footer
# ─────────────────────────────────────────
st.markdown("---")
st.caption("JM Financial · Internal Research Automation · Powered by Mistral AI · For internal use only")
